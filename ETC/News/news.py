import requests
from bs4 import BeautifulSoup
import time
import pandas as pd
import re
import sqlite3
from datetime import datetime
import hashlib
import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, simpledialog
from threading import Thread
from urllib.parse import quote
import webbrowser
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import anthropic
import os
import json
from langchain import LLMChain
from langchain.prompts import PromptTemplate
from langchain_community.llms import Anthropic
import sys
sys.stdout.reconfigure(encoding='utf-8')


# Configuration class for app settings
class Config:
    def __init__(self):
        self.config_file = "news_tool_config.json"
        self.default_config = {
            "email": {
                "smtp_server": "smtp.gmail.com",
                "smtp_port": 587,
                "sender_email": "",
                "password": ""
            },
            "llm": {
                "anthropic_api_key": "",
                "model": "claude-3-haiku-20240307",
                "temperature": 0.7,
                "max_tokens": 4000
            },
            "categories": [
                "데이터베이스 시장", "데이터베이스 경쟁사", 
                "티베로", "큐브리드", "비트나인", "인젠트",
                "오라클", "포스트그레스", "마리아디비", 
                "레디스", "몽고디비", "다이나모디비",
                "데이터베이스 클라우드", "RAG", "AI ready data", "LLM", "sLLM"
            ],
            "news_templates": {
                "standard": "# {Title}\n\n## 개요\n{Summary}\n\n## 주요 내용\n{Content}\n\n## 시장 영향\n{Impact}\n\n## 출처\n{Source}",
                "analysis": "# {Title}\n\n## 배경\n{Background}\n\n## 분석\n{Analysis}\n\n## 전망\n{Forecast}\n\n## 권장 사항\n{Recommendations}\n\n## 출처\n{Source}"
            }
        }
        self.load_config()
    
    def load_config(self):
        """설정 파일 로드"""
        try:
            if os.path.exists(self.config_file):
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    self.config = json.load(f)
            else:
                self.config = self.default_config
                self.save_config()
        except Exception as e:
            print(f"설정 로드 중 오류: {e}")
            self.config = self.default_config
    
    def save_config(self):
        """설정 파일 저장"""
        try:
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(self.config, f, indent=4, ensure_ascii=False)
        except Exception as e:
            print(f"설정 저장 중 오류: {e}")
    
    def get_categories(self):
        """뉴스 카테고리 목록 반환"""
        return self.config["categories"]
    
    def set_categories(self, categories):
        """뉴스 카테고리 목록 설정"""
        self.config["categories"] = categories
        self.save_config()
    
    def get_email_config(self):
        """이메일 설정 반환"""
        return self.config["email"]
    
    def set_email_config(self, email_config):
        """이메일 설정 저장"""
        self.config["email"] = email_config
        self.save_config()
    
    def get_llm_config(self):
        """LLM 설정 반환"""
        return self.config["llm"]
    
    def set_llm_config(self, llm_config):
        """LLM 설정 저장"""
        self.config["llm"] = llm_config
        self.save_config()
    
    def get_news_templates(self):
        """뉴스 템플릿 반환"""
        return self.config["news_templates"]
    
    def set_news_templates(self, templates):
        """뉴스 템플릿 저장"""
        self.config["news_templates"] = templates
        self.save_config()


def setup_database():
    """데이터베이스 설정 및 필요한 테이블 생성"""
    conn = sqlite3.connect('news_database.db')
    cursor = conn.cursor()
    
    # 기존 테이블 존재 여부 확인
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='news'")
    table_exists = cursor.fetchone()
    
    if table_exists:
        # 테이블 있을 경우, 컬럼 구조 확인
        cursor.execute("PRAGMA table_info(news)")
        columns = [column[1] for column in cursor.fetchall()]
        
        # search_count 컬럼이 없다면 추가
        if 'search_count' not in columns:
            cursor.execute("ALTER TABLE news ADD COLUMN search_count INTEGER DEFAULT 1")
        
        # type 컬럼이 없다면 추가 (웹 검색, AI 생성 등 구분)
        if 'type' not in columns:
            cursor.execute("ALTER TABLE news ADD COLUMN type TEXT DEFAULT 'web'")
    else:
        # 뉴스 테이블 생성 (type 컬럼 추가)
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS news (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            url TEXT UNIQUE,
            content TEXT,
            summary TEXT,
            publish_date TEXT,
            source TEXT,
            added_date TEXT NOT NULL,
            content_hash TEXT UNIQUE,
            search_count INTEGER DEFAULT 1,
            type TEXT DEFAULT 'web'
        )
        ''')
    
    # 키워드 테이블 생성
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS keywords (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        keyword TEXT UNIQUE NOT NULL
    )
    ''')
    
    # 뉴스-키워드 관계 테이블 생성
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS news_keywords (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        news_id INTEGER,
        keyword_id INTEGER,
        FOREIGN KEY (news_id) REFERENCES news (id),
        FOREIGN KEY (keyword_id) REFERENCES keywords (id),
        UNIQUE(news_id, keyword_id)
    )
    ''')
    
    # 카테고리 테이블 생성
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS categories (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT UNIQUE NOT NULL,
        description TEXT
    )
    ''')
    
    # 기본 카테고리 추가
    config = Config()
    categories = config.get_categories()
    for category in categories:
        cursor.execute("INSERT OR IGNORE INTO categories (name) VALUES (?)", (category,))
    
    conn.commit()
    return conn


def get_tables(conn):
    """데이터베이스의 테이블 목록 가져오기"""
    cursor = conn.cursor()
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
    return [table[0] for table in cursor.fetchall()]


def is_duplicate(conn, url=None, content_hash=None):
    """URL 또는 콘텐츠 해시를 기준으로 중복 여부 확인 및 검색 횟수 증가"""
    cursor = conn.cursor()
    
    if url and url.strip():
        cursor.execute("SELECT id FROM news WHERE url = ?", (url,))
        result = cursor.fetchone()
        if result:
            # 검색 횟수 증가
            cursor.execute("UPDATE news SET search_count = search_count + 1 WHERE id = ?", (result[0],))
            conn.commit()
            return True
    
    if content_hash:
        cursor.execute("SELECT id FROM news WHERE content_hash = ?", (content_hash,))
        result = cursor.fetchone()
        if result:
            # 검색 횟수 증가
            cursor.execute("UPDATE news SET search_count = search_count + 1 WHERE id = ?", (result[0],))
            conn.commit()
            return True
    
    return False


def save_to_database(conn, news_item):
    """뉴스 항목을 데이터베이스에 저장"""
    cursor = conn.cursor()
    
    # URL이 없는 경우(AI 생성 등) content_hash 생성을 위한 처리
    url = news_item.get('url', '')
    if not url:
        url = f"generated_{datetime.now().strftime('%Y%m%d%H%M%S')}"
    
    # 콘텐츠 해시 생성 (URL과 제목 기반)
    content_hash = hashlib.md5((url + news_item['title']).encode()).hexdigest()
    
    # 중복 검사
    if is_duplicate(conn, url=url if url != "" else None, content_hash=content_hash):
        print(f"중복 기사 건너뜀 (검색 횟수 증가): {news_item['title']}")
        return None
    
    # 현재 시간
    current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    
    try:
        # 뉴스 타입 결정 (기본값은 'web')
        news_type = news_item.get('type', 'web')
        
        # 뉴스 정보 저장
        cursor.execute('''
        INSERT INTO news (title, url, content, summary, publish_date, source, added_date, content_hash, search_count, type)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, 1, ?)
        ''', (
            news_item['title'],
            url,
            news_item.get('content', ''),
            news_item.get('summary', ''),
            news_item.get('publish_date', ''),
            news_item.get('source', ''),
            current_time,
            content_hash,
            news_type
        ))
        
        news_id = cursor.lastrowid
        
        # 키워드 저장
        if 'keyword' in news_item:
            # 키워드가 존재하는지 확인
            cursor.execute("SELECT id FROM keywords WHERE keyword = ?", (news_item['keyword'],))
            keyword_result = cursor.fetchone()
            
            if keyword_result:
                keyword_id = keyword_result[0]
            else:
                cursor.execute("INSERT INTO keywords (keyword) VALUES (?)", (news_item['keyword'],))
                keyword_id = cursor.lastrowid
            
            # 뉴스-키워드 관계 저장
            cursor.execute('''
            INSERT OR IGNORE INTO news_keywords (news_id, keyword_id)
            VALUES (?, ?)
            ''', (news_id, keyword_id))
        
        conn.commit()
        return news_id
    
    except sqlite3.Error as e:
        print(f"데이터베이스 저장 오류: {e}")
        conn.rollback()
        return None


def extract_publish_date(article_html):
    """기사 HTML에서 발행일을 추출하는 함수"""
    soup = BeautifulSoup(article_html, 'html.parser')
    
    # 날짜 추출을 위한 다양한 선택자
    date_selectors = [
        '.article_info .date', # 네이버 뉴스
        '.info_datestamp .date', # 다음 뉴스
        '.article_date', # 일반적인 선택자
        '.time', # 일반적인 선택자
        '.publish-date', # 일반적인 선택자
        '.date_time', # 일반적인 선택자
        'time' # HTML5 time 태그
    ]
    
    for selector in date_selectors:
        date_elem = soup.select_one(selector)
        if date_elem:
            date_text = date_elem.get_text(strip=True)
            # 날짜 텍스트 정제
            date_text = re.sub(r'입력|수정|발행|등록일?', '', date_text).strip()
            return date_text
    
    # 메타 태그에서 날짜 추출 시도
    meta_date = soup.find('meta', {'property': 'article:published_time'})
    if meta_date:
        return meta_date.get('content', '')
    
    return ''


def extract_source(article_html, url):
    """기사 HTML과 URL에서 출처를 추출하는 함수"""
    # URL에서 도메인 추출
    domain_match = re.search(r'https?://(?:www\.)?([^/]+)', url)
    if domain_match:
        domain = domain_match.group(1)
        return domain
    
    # HTML에서 언론사 정보 추출
    soup = BeautifulSoup(article_html, 'html.parser')
    
    source_selectors = [
        '.press_logo', # 네이버 뉴스
        '.article_company', # 다음 뉴스
        '.article_source', # 일반적인 선택자
        '.source', # 일반적인 선택자
        '.publisher' # 일반적인 선택자
    ]
    
    for selector in source_selectors:
        source_elem = soup.select_one(selector)
        if source_elem:
            return source_elem.get_text(strip=True)
    
    return ''


def search_naver_news(keywords, conn, status_callback=None, result_callback=None):
    """
    네이버 뉴스에서 키워드를 검색하고 결과를 반환하는 함수
    """
    all_results = []
    
    # 각 키워드에 대해 검색 실행
    for keyword in keywords:
        if status_callback:
            status_callback(f"검색 중: {keyword}")
        else:
            print(f"검색 중: {keyword}")
        
        # 네이버 뉴스 검색 URL (검색어 URL 인코딩)
        encoded_keyword = quote(keyword)
        search_url = f'https://search.naver.com/search.naver?where=news&sm=tab_jum&query={encoded_keyword}'
        
        # 요청 헤더 설정
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        # 검색 페이지 요청
        try:
            response = requests.get(search_url, headers=headers)
            
            if response.status_code == 200:
                soup = BeautifulSoup(response.text, 'html.parser')
                
                # 네이버 뉴스 결과 컨테이너 찾기
                news_items = soup.select('.bx')
                
                if not news_items:
                    news_items = soup.select('.news_area')  # 다른 클래스 시도
                
                for item in news_items:
                    try:
                        # 뉴스 제목과 링크 추출 - 여러 선택자 시도
                        title_elem = None
                        
                        for selector in ['.news_tit', '.title_link']:
                            title_elem = item.select_one(selector)
                            if title_elem:
                                break
                        
                        if title_elem:
                            title = title_elem.get_text(strip=True)
                            url = title_elem.get('href')
                            
                            # 미리보기 텍스트 추출 - 여러 선택자 시도
                            preview_text = ""
                            
                            for selector in ['.dsc_txt', '.api_txt_lines', '.desc']:
                                preview_elem = item.select_one(selector)
                                if preview_elem:
                                    preview_text = preview_elem.get_text(strip=True)
                                    break
                            
                            # 중복 검사 - 검색 횟수 증가 포함
                            duplicate = is_duplicate(conn, url=url)
                            if duplicate:
                                if status_callback:
                                    status_callback(f"  중복 기사 건너뜀 (검색 횟수 증가): {title}")
                                else:
                                    print(f"  중복 기사 건너뜀 (검색 횟수 증가): {title}")
                                
                                # 중복이라도 결과로 반환하기 위해 DB에서 기사 정보 가져오기
                                cursor = conn.cursor()
                                cursor.execute('''
                                SELECT n.id, n.title, n.url, n.summary, n.content, n.publish_date, n.source, 
                                       n.added_date, n.search_count, k.keyword, n.type
                                FROM news n
                                JOIN news_keywords nk ON n.id = nk.news_id
                                JOIN keywords k ON nk.keyword_id = k.id
                                WHERE n.url = ?
                                ''', (url,))
                                
                                row = cursor.fetchone()
                                if row:
                                    news_item = {
                                        'id': row[0],
                                        'title': row[1],
                                        'url': row[2],
                                        'summary': row[3],
                                        'content': row[4],
                                        'publish_date': row[5],
                                        'source': row[6],
                                        'added_date': row[7],
                                        'search_count': row[8],
                                        'keyword': row[9],
                                        'type': row[10],
                                        'is_duplicate': True
                                    }
                                    all_results.append(news_item)
                                    
                                    # 검색 결과 콜백 호출
                                    if result_callback:
                                        result_callback(news_item)
                                
                                continue
                            
                            # 결과 저장
                            news_item = {
                                'keyword': keyword,
                                'title': title,
                                'url': url,
                                'preview': preview_text,
                                'is_duplicate': False,
                                'search_count': 1,
                                'type': 'web'
                            }
                            
                            # 기사 본문과 요약 가져오기
                            if status_callback:
                                status_callback(f"  기사 내용 가져오는 중: {title}")
                            
                            content, html = get_article_content(url)
                            if content:
                                news_item['content'] = content
                                
                                # 미리보기가 충분히 길면 그대로 사용
                                if len(preview_text) > 150:
                                    summary = preview_text
                                else:
                                    # 그렇지 않으면 본문에서 요약 생성
                                    summary = generate_summary(content)
                                
                                news_item['summary'] = summary
                                
                                # 발행일과 출처 추출
                                news_item['publish_date'] = extract_publish_date(html)
                                news_item['source'] = extract_source(html, url)
                                
                                # 데이터베이스에 저장
                                news_id = save_to_database(conn, news_item)
                                if news_id:
                                    news_item['id'] = news_id
                                    all_results.append(news_item)
                                    
                                    if status_callback:
                                        status_callback(f"  저장 완료: {title}")
                                    else:
                                        print(f"  저장 완료: {title}")
                                    
                                    # 검색 결과 콜백 호출
                                    if result_callback:
                                        result_callback(news_item)
                                else:
                                    if status_callback:
                                        status_callback(f"  저장 실패: {title}")
                                    else:
                                        print(f"  저장 실패: {title}")
                            else:
                                if status_callback:
                                    status_callback(f"  본문을 가져올 수 없음: {title}")
                                else:
                                    print(f"  본문을 가져올 수 없음: {title}")
                                
                    except Exception as e:
                        if status_callback:
                            status_callback(f"항목 처리 중 오류 발생: {e}")
                        else:
                            print(f"항목 처리 중 오류 발생: {e}")
                
                # 과도한 요청 방지를 위한 딜레이
                time.sleep(1)
            else:
                if status_callback:
                    status_callback(f"검색 실패 (상태 코드: {response.status_code})")
                else:
                    print(f"검색 실패 (상태 코드: {response.status_code})")
        except Exception as e:
            if status_callback:
                status_callback(f"요청 중 오류 발생: {e}")
            else:
                print(f"요청 중 오류 발생: {e}")
    
    return all_results


def get_article_content(url):
    """뉴스 기사의 본문 내용을 가져오는 함수"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(url, headers=headers)
        if response.status_code == 200:
            html = response.text
            soup = BeautifulSoup(html, 'html.parser')
            
            # 여러 뉴스 사이트의 본문 컨테이너 선택자
            content_selectors = [
                '#articleBodyContents', # 네이버 뉴스
                '.article_view', # 다음 뉴스
                '#article-view-content-div', # 많은 언론사
                '.news_content', # 일부 뉴스사이트
                '.article-body', # 일부 뉴스사이트
                '.article_body', # 일부 뉴스사이트
                '.newsct_article', # 일부 뉴스사이트
                '#newsEndContents', # 연합뉴스
                '.news_body', # 일부 사이트
                '.article-text', # 일부 사이트
                '#newsContent' # 일부 사이트
            ]
            
            content_text = ""
            
            for selector in content_selectors:
                content_container = soup.select_one(selector)
                if content_container:
                    # 불필요한 요소 제거
                    for unwanted in content_container.select('script, style, iframe, .reporter_area, .byline'):
                        unwanted.extract()
                    
                    # 텍스트 가져오기
                    content_text = content_container.get_text('\n', strip=True)
                    
                    # 본문이 충분히 길면 중단
                    if len(content_text) > 200:
                        break
            
            # 본문을 찾지 못한 경우 일반 p 태그에서 추출 시도
            if not content_text:
                paragraphs = soup.select('p')
                content_text = '\n'.join([p.get_text(strip=True) for p in paragraphs if len(p.get_text(strip=True)) > 50])
            
            # 여러 줄의 공백 제거
            content_text = re.sub(r'\n\s*\n', '\n', content_text)
            
            # 내용이 충분하지 않으면 빈 문자열 반환
            if len(content_text) < 100:
                return "", html
                
            return content_text, html
        else:
            return "", ""
    except Exception as e:
        print(f"기사 내용 가져오기 오류: {e}")
        return "", ""


def generate_summary(content, max_lines=5):
    """기사 내용에서 요약 생성 (첫 5줄 추출)"""
    if not content:
        return "요약을 가져올 수 없습니다."
    
    # 텍스트를 줄 단위로 분리
    lines = [line for line in content.split('\n') if line.strip()]
    
    # 첫 max_lines 줄 선택 (또는 모든 줄이 그보다 적으면 모든 줄)
    summary_lines = lines[:max_lines]
    summary = '\n'.join(summary_lines)
    
    return summary


def save_to_word(results, filename="news_results.docx"):
    """검색 결과를 Word 문서로 저장하는 함수"""
    doc = Document()
    
    # 제목 추가
    title = doc.add_heading('데이터베이스 관련 뉴스 검색 결과', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 현재 날짜 추가
    current_date = datetime.now().strftime('%Y년 %m월 %d일')
    date_paragraph = doc.add_paragraph(f'작성일: {current_date}')
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 소개 텍스트 추가
    doc.add_paragraph('아래와 같이 주요 데이터베이스 관련 뉴스를 공유드립니다. 업무에 도움이 되시길 바랍니다.')
    doc.add_paragraph('')
    
    # 각 검색 결과 추가
    for result in results:
        # 뉴스 제목
        heading = doc.add_heading(level=2)
        heading_run = heading.add_run(f"{result['title']}")
        heading_run.bold = True
        
        # 출처 및 날짜 추가
        source_date = ""
        if 'source' in result and result['source']:
            source_date += result['source']
        if 'publish_date' in result and result['publish_date']:
            if source_date:
                source_date += " | "
            source_date += result['publish_date']
        
        if source_date:
            source_para = doc.add_paragraph(source_date)
            source_para.italic = True
        
        # 링크 추가 (AI 생성 뉴스는 링크가 없을 수 있음)
        if 'url' in result and result['url'] and not result['url'].startswith('generated_'):
            link_para = doc.add_paragraph('링크: ')
            link_para.add_run(result['url'])
        
        # 키워드 추가
        keyword_para = doc.add_paragraph('키워드: ')
        keyword_para.add_run(result['keyword']).bold = True
        
        # 검색 횟수 추가
        search_count_para = doc.add_paragraph('검색 횟수: ')
        search_count_para.add_run(str(result.get('search_count', 1))).bold = True
        
        # 타입 추가 (웹 또는 AI 생성)
        type_para = doc.add_paragraph('타입: ')
        type_para.add_run(result.get('type', 'web')).italic = True
        
        # 요약 추가
        doc.add_paragraph('요약:')
        doc.add_paragraph(result.get('summary', '요약 없음'))
        
        # 구분선 추가
        doc.add_paragraph('---------------------------------------------------')
    
    # 문서 저장
    doc.save(filename)
    return f"검색 결과가 '{filename}' 파일로 저장되었습니다."


def generate_markdown(results, filename="news_results.md"):
    """검색 결과를 마크다운 파일로 저장하는 함수"""
    with open(filename, 'w', encoding='utf-8') as f:
        f.write("# 뉴스 검색 결과\n\n")
        f.write(f"*작성일: {datetime.now().strftime('%Y년 %m월 %d일')}*\n\n")
        
        for result in results:
            if 'url' in result and result['url'] and not result['url'].startswith('generated_'):
                f.write(f"## [{result['title']}]({result['url']})\n\n")
            else:
                f.write(f"## {result['title']}\n\n")
            
            # 출처 및 날짜 추가
            source_date = ""
            if 'source' in result and result['source']:
                source_date += result['source']
            if 'publish_date' in result and result['publish_date']:
                if source_date:
                    source_date += " | "
                source_date += result['publish_date']
            
            if source_date:
                f.write(f"*{source_date}*\n\n")
            
            f.write(f"**키워드**: {result['keyword']}\n\n")
            f.write(f"**검색 횟수**: {result.get('search_count', 1)}\n\n")
            f.write(f"**타입**: {result.get('type', 'web')}\n\n")
            f.write("**요약**:\n")
            f.write(result.get('summary', '요약 없음') + "\n\n")
            f.write("---\n\n")
    
    return f"검색 결과가 '{filename}' 파일로 저장되었습니다."


class LLMNewsGenerator:
    """LLM을 활용한 뉴스 생성 클래스"""
    def __init__(self):
        self.config = Config()
        self.llm_config = self.config.get_llm_config()
        self.templates = self.config.get_news_templates()
    
    def get_anthropic_client(self):
        """Anthropic 클라이언트 반환"""
        try:
            api_key = self.llm_config.get("anthropic_api_key", "")
            if not api_key:
                return None
            
            client = anthropic.Anthropic(api_key=api_key)
            return client
        except Exception as e:
            print(f"Anthropic 클라이언트 생성 오류: {e}")
            return None
    
    def generate_news(self, input_data, template_key="standard", page_length=1):
        """뉴스 생성"""
        client = self.get_anthropic_client()
        if not client:
            return {"error": "API 키가 설정되지 않았습니다. 설정 메뉴에서 LLM API 키를 설정해주세요."}
        
        # 템플릿 선택
        template = self.templates.get(template_key, self.templates["standard"])
        
        # 페이지 길이에 따른 토큰 설정
        max_tokens = {
            1: 4000,
            2: 8000,
            4: 15000
        }.get(page_length, 4000)
        
        # 프롬프트 생성
        template_keys = re.findall(r'\{([^}]+)\}', template)
        
        system_prompt = f"""
        당신은 전문적인 데이터베이스 및 IT 기술 뉴스 작성자입니다.
        사용자가 제공한 정보를 바탕으로 {page_length}페이지 분량의 전문적인 뉴스/기술 보고서를 작성해주세요.
        
        다음 틀에 맞추어 작성하세요:
        {template}
        
        작성 시 주의사항:
        1. 제공된 정보를 바탕으로 전문적이고 명확한 내용으로 작성하세요.
        2. 주요 기술 용어와 개념을 정확하게 설명하고 관련 산업 동향을 포함하세요.
        3. 중립적이고 객관적인 톤을 유지하되, 실용적인 분석과 통찰을 제공하세요.
        4. 문장은 간결하고 정확하게 작성하며, 기술 용어는 필요시 간단한 설명을 덧붙이세요.
        5. 뉴스의 분량은 약 {page_length}페이지 정도로 작성하세요.
        6. 주요 내용, 분석, 결론을 포함하는 구조로 작성하세요.
        """
        
        # 사용자 메시지 생성
        user_prompt = "다음 정보를 바탕으로 뉴스/보고서를 작성해주세요:\n\n"
        
        for key in template_keys:
            if key in input_data:
                user_prompt += f"[{key}]\n{input_data[key]}\n\n"
        
        try:
            # Anthropic API 호출
            response = client.messages.create(
                model=self.llm_config.get("model", "claude-3-haiku-20240307"),
                max_tokens=max_tokens,
                temperature=self.llm_config.get("temperature", 0.7),
                system=system_prompt,
                messages=[
                    {"role": "user", "content": user_prompt}
                ]
            )
            
            # 결과 반환
            return {"content": response.content[0].text, "error": None}
        
        except Exception as e:
            print(f"뉴스 생성 오류: {e}")
            return {"content": None, "error": f"뉴스 생성 중 오류가 발생했습니다: {e}"}
    
    def search_related_articles(self, conn, keyword, count=5):
        """관련 뉴스 기사 검색"""
        cursor = conn.cursor()
        cursor.execute("""
        SELECT n.id, n.title, n.summary, n.content, k.keyword
        FROM news n
        JOIN news_keywords nk ON n.id = nk.news_id
        JOIN keywords k ON nk.keyword_id = k.id
        WHERE k.keyword = ? OR n.title LIKE ? OR n.content LIKE ?
        ORDER BY n.added_date DESC
        LIMIT ?
        """, (keyword, f"%{keyword}%", f"%{keyword}%", count))
        
        results = cursor.fetchall()
        
        related_articles = []
        for row in results:
            related_articles.append({
                'id': row[0],
                'title': row[1],
                'summary': row[2],
                'content': row[3],
                'keyword': row[4]
            })
        
        return related_articles
    
    def generate_with_rag(self, conn, input_data, template_key="standard", page_length=1):
        """RAG 기술을 활용한 뉴스 생성"""
        # 관련 기사 검색
        keyword = input_data.get("keyword", "")
        if not keyword:
            return {"error": "키워드가 제공되지 않았습니다."}
        
        related_articles = self.search_related_articles(conn, keyword)
        
        # 관련 자료 컨텍스트 생성
        context = "## 관련 기존 기사 정보\n\n"
        for article in related_articles:
            context += f"### {article['title']}\n"
            context += f"요약: {article['summary']}\n\n"
        
        # 입력 데이터에 컨텍스트 추가
        enhanced_input = input_data.copy()
        if "Content" in enhanced_input:
            enhanced_input["Content"] += f"\n\n{context}"
        else:
            enhanced_input["Content"] = context
        
        # LLM 호출
        return self.generate_news(enhanced_input, template_key, page_length)
    
class SettingsDialog:
    def __init__(self, parent):
        self.parent = parent
        self.config = Config()
        
        # 대화상자 생성
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("프로그램 설정")
        self.dialog.geometry("500x600")
        self.dialog.resizable(False, False)
        self.dialog.transient(parent)
        self.dialog.grab_set()
        
        # 탭 컨트롤
        self.notebook = ttk.Notebook(self.dialog)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 이메일 설정 탭
        self.email_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.email_frame, text="이메일 설정")
        self.create_email_settings()
        
        # LLM 설정 탭
        self.llm_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.llm_frame, text="LLM 설정")
        self.create_llm_settings()
        
        # 카테고리 설정 탭
        self.categories_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.categories_frame, text="카테고리 설정")
        self.create_categories_settings()
        
        # 템플릿 설정 탭
        self.templates_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.templates_frame, text="템플릿 설정")
        self.create_templates_settings()
        
        # 저장 버튼
        save_button = ttk.Button(self.dialog, text="저장", command=self.save_settings)
        save_button.pack(side=tk.RIGHT, padx=10, pady=10)
        
        # 취소 버튼
        cancel_button = ttk.Button(self.dialog, text="취소", command=self.dialog.destroy)
        cancel_button.pack(side=tk.RIGHT, padx=10, pady=10)

    def add_category(self):
        """카테고리 추가"""
        new_category = self.new_category_var.get().strip()
        if new_category:
            # 현재 카테고리 목록 가져오기
            current_categories = self.categories_text.get(1.0, tk.END).strip().split('\n')
            
            # 중복 체크
            if new_category in current_categories:
                messagebox.showwarning("경고", f"카테고리 '{new_category}'은(는) 이미 존재합니다.")
                return
            
            # 텍스트 영역에 추가
            self.categories_text.insert(tk.END, new_category + "\n")
            self.new_category_var.set("")

    def test_email(self):
        """이메일 설정 테스트"""
        # 현재 입력된 값으로 임시 설정
        email_config = {
            "smtp_server": self.smtp_server_var.get(),
            "smtp_port": self.smtp_port_var.get(),
            "sender_email": self.sender_email_var.get(),
            "password": self.password_var.get()
        }
        
        # 필수 값 확인
        if not all([email_config["smtp_server"], email_config["sender_email"], email_config["password"]]):
            messagebox.showerror("오류", "모든 이메일 설정 값을 입력해주세요.")
            return
        
        try:
            # 테스트 이메일 주소 입력
            test_email = simpledialog.askstring("테스트", "테스트 이메일 주소를 입력하세요:")
            if not test_email:
                return
            
            # 테스트 이메일 보내기
            with smtplib.SMTP(email_config["smtp_server"], email_config["smtp_port"]) as server:
                server.starttls()
                server.login(email_config["sender_email"], email_config["password"])
                
                msg = MIMEMultipart()
                msg['From'] = email_config["sender_email"]
                msg['To'] = test_email
                msg['Subject'] = "뉴스 검색 도구 - 이메일 테스트"
                
                body = "이 메일은 뉴스 검색 도구의 이메일 설정 테스트입니다."
                msg.attach(MIMEText(body, 'plain'))
                
                server.send_message(msg)
            
            messagebox.showinfo("성공", f"테스트 이메일을 {test_email}로 보냈습니다.")
        except Exception as e:
            messagebox.showerror("오류", f"이메일 테스트 실패: {e}")

    def test_llm(self):
        """LLM API 설정 테스트"""
        # 현재 입력된 값으로 임시 설정
        llm_config = {
            "anthropic_api_key": self.api_key_var.get(),
            "model": self.model_var.get(),
            "temperature": self.temperature_var.get()
        }
        
        # API 키 확인
        if not llm_config["anthropic_api_key"]:
            messagebox.showerror("오류", "Anthropic API 키를 입력해주세요.")
            return
        
        try:
            # Anthropic 클라이언트 생성
            client = anthropic.Anthropic(api_key=llm_config["anthropic_api_key"])
            
            # 간단한 테스트 메시지
            response = client.messages.create(
                model=llm_config["model"],
                max_tokens=100,
                temperature=llm_config["temperature"],
                messages=[
                    {"role": "user", "content": "데이터베이스 기술에 대해 한 문장으로 설명해주세요."}
                ]
            )
            
            messagebox.showinfo("성공", f"API 테스트 성공:\n\n{response.content[0].text}")
        except Exception as e:
            messagebox.showerror("오류", f"API 테스트 실패: {e}")

    def save_settings(self):
        """설정 저장"""
        try:
            # 이메일 설정 저장
            email_config = {
                "smtp_server": self.smtp_server_var.get(),
                "smtp_port": self.smtp_port_var.get(),
                "sender_email": self.sender_email_var.get(),
                "password": self.password_var.get()
            }
            self.config.set_email_config(email_config)
            
            # LLM 설정 저장
            llm_config = {
                "anthropic_api_key": self.api_key_var.get(),
                "model": self.model_var.get(),
                "temperature": self.temperature_var.get(),
                "max_tokens": self.config.get_llm_config()["max_tokens"]  # 기존 값 유지
            }
            self.config.set_llm_config(llm_config)
            
            # 카테고리 설정 저장
            categories = self.categories_text.get(1.0, tk.END).strip().split('\n')
            categories = [cat for cat in categories if cat.strip()]  # 빈 줄 제거
            self.config.set_categories(categories)
            
            # 현재 선택된 템플릿 저장
            self.save_current_template()
            
            messagebox.showinfo("알림", "모든 설정이 저장되었습니다.")
            self.dialog.destroy()
            
            # 부모 창에 설정 변경 알림
            if hasattr(self.parent, 'refresh_categories'):
                self.parent.refresh_categories()
        
        except Exception as e:
            messagebox.showerror("오류", f"설정 저장 중 오류 발생: {e}")

# GUI 클래스
class NewsSearchApp:
    def __init__(self, root):
        self.root = root
        self.root.title("데이터베이스 뉴스 검색 및 분석 도구")
        self.root.geometry("1200x800")
        
        # 데이터베이스 연결
        self.conn = setup_database()
        
        # 설정 로드
        self.config = Config()
        
        # 저장된 결과
        self.search_results = []
        
        self.create_widgets()
    
    def create_widgets(self):
        # 메인 프레임
        main_frame = ttk.Frame(self.root, padding=10)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 메뉴바 생성
        self.create_menu()
        
        # 상단 컨트롤 프레임
        control_frame = ttk.LabelFrame(main_frame, text="검색 옵션", padding=10)
        control_frame.pack(fill=tk.X, pady=5)
        
        # 키워드 선택 영역
        keyword_frame = ttk.Frame(control_frame)
        keyword_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(keyword_frame, text="키워드:").pack(side=tk.LEFT, padx=5)
        
        self.keyword_var = tk.StringVar()
        self.keyword_combobox = ttk.Combobox(keyword_frame, textvariable=self.keyword_var, width=30)
       
        self.keyword_combobox.pack(side=tk.LEFT, padx=5)
        
        # 추가 필터 영역
        filter_frame = ttk.Frame(control_frame)
        filter_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(filter_frame, text="검색어:").pack(side=tk.LEFT, padx=5)
        self.search_term_entry = ttk.Entry(filter_frame, width=30)
        self.search_term_entry.pack(side=tk.LEFT, padx=5)
        
        # 뉴스 타입 필터
        ttk.Label(filter_frame, text="타입:").pack(side=tk.LEFT, padx=(10, 5))
        self.news_type_var = tk.StringVar(value="all")
        ttk.Radiobutton(filter_frame, text="전체", variable=self.news_type_var, value="all").pack(side=tk.LEFT)
        ttk.Radiobutton(filter_frame, text="웹 검색", variable=self.news_type_var, value="web").pack(side=tk.LEFT)
        ttk.Radiobutton(filter_frame, text="AI 생성", variable=self.news_type_var, value="ai").pack(side=tk.LEFT)
        
        # 버튼 영역
        button_frame = ttk.Frame(control_frame)
        button_frame.pack(fill=tk.X, pady=10)
        
        self.search_button = ttk.Button(button_frame, text="검색", command=self.start_search)
        self.search_button.pack(side=tk.LEFT, padx=5)
        
        
    
        
        self.refresh_button = ttk.Button(button_frame, text="DB 새로고침", command=self.refresh_db_data)
        self.refresh_button.pack(side=tk.LEFT, padx=5)
        
        # 상태 표시 영역
        status_frame = ttk.Frame(main_frame)
        status_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(status_frame, text="상태:").pack(side=tk.LEFT, padx=5)
        self.status_var = tk.StringVar(value="준비됨")
        ttk.Label(status_frame, textvariable=self.status_var).pack(side=tk.LEFT, padx=5)
        
        # 결과 테이블
        table_frame = ttk.LabelFrame(main_frame, text="검색 결과", padding=10)
        table_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # 테이블 스타일 설정
        style = ttk.Style()
        style.configure("Treeview", font=('Arial', 10))
        style.configure("Treeview.Heading", font=('Arial', 10, 'bold'))
        
        # 스크롤바 설정
        table_scroll = ttk.Scrollbar(table_frame)
        table_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 테이블 생성
        self.result_table = ttk.Treeview(table_frame, 
                                       columns=("id", "title", "source", "date", "keyword", "count", "type"),
                                       show="headings",
                                       yscrollcommand=table_scroll.set)
        
        # 스크롤바 연결
        table_scroll.config(command=self.result_table.yview)
        
        # 컬럼 설정
        self.result_table.heading("id", text="ID")
        self.result_table.heading("title", text="제목")
        self.result_table.heading("source", text="출처")
        self.result_table.heading("date", text="날짜")
        self.result_table.heading("keyword", text="키워드")
        self.result_table.heading("count", text="검색횟수")
        self.result_table.heading("type", text="타입")
        
        self.result_table.column("id", width=50, anchor=tk.CENTER)
        self.result_table.column("title", width=350)
        self.result_table.column("source", width=120)
        self.result_table.column("date", width=120)
        self.result_table.column("keyword", width=120)
        self.result_table.column("count", width=80, anchor=tk.CENTER)
        self.result_table.column("type", width=80, anchor=tk.CENTER)
        
        self.result_table.pack(fill=tk.BOTH, expand=True)

        
        # 상세 정보 영역
        detail_frame = ttk.LabelFrame(main_frame, text="기사 상세 정보", padding=10)
        detail_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        self.detail_text = scrolledtext.ScrolledText(detail_frame, wrap=tk.WORD, height=10)
        self.detail_text.pack(fill=tk.BOTH, expand=True)
        
        # 초기 데이터 로드
        self.refresh_db_data()

    def refresh_db_data(self):
        """데이터베이스에서 데이터 새로고침"""
        # 테이블 초기화
        for item in self.result_table.get_children():
            self.result_table.delete(item)
        
        # 최근 뉴스 가져오기 (최대 100개)
        cursor = self.conn.cursor()
        cursor.execute("""
            SELECT n.id, n.title, n.source, n.publish_date, k.keyword, n.search_count, n.url, n.summary, n.type
            FROM news n
            JOIN news_keywords nk ON n.id = nk.news_id
            JOIN keywords k ON nk.keyword_id = k.id
            ORDER BY n.added_date DESC
            LIMIT 100
        """)
        rows = cursor.fetchall()
        
        # 테이블에 데이터 추가
        for row in rows:
            self.result_table.insert("", tk.END, values=(row[0], row[1], row[2], row[3], row[4], row[5], row[8]),
                                    iid=row[0], tags=("clickable",))
        
        self.status_var.set(f"데이터베이스에서 {len(rows)}개의 기사를 불러왔습니다.")
        self.search_results = []
        
    def start_search(self):
        """검색 시작"""
        keyword = self.keyword_var.get()
        search_term = self.search_term_entry.get()
        news_type = self.news_type_var.get()
        
        if not keyword and not search_term and news_type == "all":
            messagebox.showwarning("경고", "키워드, 검색어 또는 뉴스 타입을 지정하세요.")
            return
        
        # 검색 옵션에 따라 다른 검색 실행
        if keyword and not search_term and news_type == "all":
            # 새 기사 검색 (웹)
            self.search_button.config(state=tk.DISABLED)
            self.status_var.set("검색 중...")
            
            # 검색 스레드 시작
            search_thread = Thread(target=self.run_search_thread, args=([keyword],))
            search_thread.daemon = True
            search_thread.start()
        else:
            # 데이터베이스 내 검색
            filter_type = None if news_type == "all" else news_type
            results = search_news_from_db(
                self.conn, 
                search_term=search_term, 
                keyword=keyword if keyword else None,
                news_type=filter_type
            )
            self.update_results_table(results)
            

    def create_menu(self):
        """메뉴바 생성"""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        # 파일 메뉴
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="파일", menu=file_menu)
        file_menu.add_command(label="검색", command=self.start_search)
        file_menu.add_command(label="새로고침", command=self.refresh_db_data)
        file_menu.add_separator()
       
        file_menu.add_separator()
        file_menu.add_command(label="종료", command=self.root.quit)
        
        # 뉴스 메뉴
        news_menu = tk.Menu(menubar, tearoff=0)
        # 설정 메뉴
        settings_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="설정", menu=settings_menu)

        
        # 도움말 메뉴
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="도움말", menu=help_menu)

        
    def send_email(to_emails, subject, body, attachments=None):
        """이메일 전송 함수"""
        config = Config()
        email_config = config.get_email_config()
        
        if not email_config["sender_email"] or not email_config["password"]:
            return "이메일 설정이 완료되지 않았습니다. 설정 메뉴에서 이메일 정보를 입력해주세요."
        
        try:
            # 이메일 객체 생성
            msg = MIMEMultipart()
            msg['From'] = email_config["sender_email"]
            msg['To'] = ", ".join(to_emails)
            msg['Subject'] = subject
            
            # 본문 추가
            msg.attach(MIMEText(body, 'html'))
            
            # 첨부파일 추가
            if attachments:
                for file_path in attachments:
                    with open(file_path, 'rb') as file:
                        part = MIMEApplication(file.read(), Name=os.path.basename(file_path))
                        part['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_path)}"'
                        msg.attach(part)
            
            # SMTP 서버 연결 및 로그인
            with smtplib.SMTP(email_config["smtp_server"], email_config["smtp_port"]) as server:
                server.starttls()
                server.login(email_config["sender_email"], email_config["password"])
                server.send_message(msg)
            
            return "이메일이 성공적으로 전송되었습니다."
        except Exception as e:
            return f"이메일 전송 중 오류가 발생했습니다: {e}"

    def create_email_content(news_items):
        """이메일 내용 생성 함수"""
        html = """
        <html>
        <head>
            <style>
                body { font-family: Arial, sans-serif; line-height: 1.6; }
                .container { width: 90%; margin: 0 auto; }
                .header { background-color: #4A86E8; color: white; padding: 10px; text-align: center; }
                .news-item { border: 1px solid #ddd; margin: 10px 0; padding: 15px; }
                .news-title { font-size: 18px; font-weight: bold; color: #2C5AA0; }
                .news-meta { color: #666; font-size: 12px; margin: 5px 0; }
                .news-summary { margin-top: 10px; }
                .footer { margin-top: 20px; font-size: 12px; color: #666; text-align: center; }
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <h2>데이터베이스 뉴스 업데이트</h2>
                </div>
                <p>안녕하세요, 아래와 같이 데이터베이스 관련 최신 뉴스를 공유드립니다.</p>
        """
        
        for item in news_items:
            html += f"""
            <div class="news-item">
                <div class="news-title">
            """
            
            if 'url' in item and item['url'] and not item['url'].startswith('generated_'):
                html += f'<a href="{item["url"]}">{item["title"]}</a>'
            else:
                html += f'{item["title"]}'
            
            html += '</div>'
            
            # 출처 및 날짜
            meta = []
            if 'source' in item and item['source']:
                meta.append(item['source'])
            if 'publish_date' in item and item['publish_date']:
                meta.append(item['publish_date'])
            
            if meta:
                html += f'<div class="news-meta">{" | ".join(meta)}</div>'
            
            # 키워드 및 타입
            html += f'<div class="news-meta">키워드: <b>{item["keyword"]}</b> | 타입: {item.get("type", "web")}</div>'
            
            # 요약
            html += f'<div class="news-summary">{item.get("summary", "요약 없음")}</div>'
            
            html += '</div>'
        
        html += """
                <div class="footer">
                    <p>본 메일은 자동 생성된 뉴스 요약입니다. 문의사항은 관리자에게 연락주세요.</p>
                </div>
            </div>
        </body>
        </html>
        """
        
        return html

    def search_news_from_db(conn, search_term=None, keyword=None, date_from=None, date_to=None, news_type=None):
        """데이터베이스에서 뉴스 검색"""
        cursor = conn.cursor()
        
        query = """
        SELECT n.id, n.title, n.url, n.summary, n.content, n.publish_date, n.source, 
            n.added_date, n.search_count, k.keyword, n.type
        FROM news n
        JOIN news_keywords nk ON n.id = nk.news_id
        JOIN keywords k ON nk.keyword_id = k.id
        WHERE 1=1
        """
        
        params = []
        
        if search_term:
            query += " AND (n.title LIKE ? OR n.content LIKE ?)"
            search_param = f"%{search_term}%"
            params.extend([search_param, search_param])
        
        if keyword:
            query += " AND k.keyword = ?"
            params.append(keyword)
        
        if date_from:
            query += " AND n.publish_date >= ?"
            params.append(date_from)
        
        if date_to:
            query += " AND n.publish_date <= ?"
            params.append(date_to)
        
        if news_type:
            query += " AND n.type = ?"
            params.append(news_type)
        
        query += " ORDER BY n.publish_date DESC"
        
        cursor.execute(query, params)
        results = cursor.fetchall()
        
        news_list = []
        for row in results:
            news_list.append({
                'id': row[0],
                'title': row[1],
                'url': row[2],
                'summary': row[3],
                'content': row[4],
                'publish_date': row[5],
                'source': row[6],
                'added_date': row[7],
                'search_count': row[8],
                'keyword': row[9],
                'type': row[10]
            })
        
        return news_list

    def get_all_keywords(conn):
        """데이터베이스에서 모든 키워드 가져오기"""
        cursor = conn.cursor()
        cursor.execute("SELECT keyword FROM keywords ORDER BY keyword")
        results = cursor.fetchall()
        return [row[0] for row in results]

    def get_news_by_id(conn, news_id):
        """ID로 뉴스 항목 가져오기"""
        cursor = conn.cursor()
        cursor.execute("""
        SELECT n.id, n.title, n.url, n.summary, n.content, n.publish_date, n.source, 
            n.added_date, n.search_count, k.keyword, n.type
        FROM news n
        LEFT JOIN news_keywords nk ON n.id = nk.news_id
        LEFT JOIN keywords k ON nk.keyword_id = k.id
        WHERE n.id = ?
        """, (news_id,))
        
        row = cursor.fetchone()
        if row:
            return {
                'id': row[0],
                'title': row[1],
                'url': row[2],
                'summary': row[3],
                'content': row[4],
                'publish_date': row[5],
                'source': row[6],
                'added_date': row[7],
                'search_count': row[8],
                'keyword': row[9],
                'type': row[10]
            }
        
        return None

def main():
    root = tk.Tk()
    app = NewsSearchApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()