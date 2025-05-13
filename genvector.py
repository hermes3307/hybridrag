#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Vector Store Generator for RAG
------------------------------
A CLI tool that processes documents, chunks them, and creates vector embeddings
for use in retrieval-augmented generation (RAG) systems.
Based on the HybridRAG framework and adapted for vectorization focus.
"""

import os
import sys
import time
import json
import argparse
import logging
import re
import traceback
from pathlib import Path
from typing import List, Dict, Tuple, Any, Optional, Set
from datetime import datetime
from collections import deque

# Document processing
import docx
import PyPDF2
import pandas as pd

# ML/Vector libraries
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

# Terminal UI
from rich.console import Console
from rich.table import Table
from rich.panel import Panel
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn, TaskID
from rich.prompt import Prompt, Confirm
from rich import print as rprint
from rich.layout import Layout
from rich.text import Text

# NLP support
import spacy

# Initialize console
console = Console()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('vector_generator.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("Vector_Generator")

class ProgressTracker:
    """Track progress with ETA calculation"""
    
    def __init__(self):
        self.start_time = time.time()
        self.file_times = deque(maxlen=50)
        self.current_file_start = None
        self.total_files = 0
        self.completed_files = 0
        self.current_phase = ""
        
    def start_tracking(self, total_files: int, phase: str = ""):
        """Start tracking progress"""
        self.total_files = total_files
        self.completed_files = 0
        self.start_time = time.time()
        self.current_phase = phase
        self.file_times.clear()
        
    def start_file(self):
        """Mark the start of file processing"""
        self.current_file_start = time.time()
    
    def complete_file(self, file_name: str = None):
        """Mark file as completed and update statistics"""
        if self.current_file_start:
            processing_time = time.time() - self.current_file_start
            self.file_times.append(processing_time)
            self.current_file_start = None
        
        self.completed_files += 1
        return self.get_progress_info(file_name)
    
    def get_progress_info(self, current_file: str = None) -> Dict[str, Any]:
        """Get current progress information"""
        if self.total_files == 0:
            return {}
        
        percentage = (self.completed_files / self.total_files) * 100
        elapsed_time = time.time() - self.start_time
        
        # Calculate ETA
        eta_seconds = None
        avg_time_per_file = None
        
        if self.file_times and self.completed_files > 0:
            avg_time_per_file = sum(self.file_times) / len(self.file_times)
            remaining_files = self.total_files - self.completed_files
            if remaining_files > 0:
                eta_seconds = avg_time_per_file * remaining_files
        
        # Format times
        elapsed_str = self._format_time(elapsed_time)
        eta_str = self._format_time(eta_seconds) if eta_seconds else "calculating..."
        avg_time_str = self._format_time(avg_time_per_file) if avg_time_per_file else "calculating..."
        
        return {
            'current_file': current_file,
            'completed': self.completed_files,
            'total': self.total_files,
            'percentage': percentage,
            'elapsed_time': elapsed_str,
            'eta': eta_str,
            'avg_time_per_file': avg_time_str,
            'remaining': self.total_files - self.completed_files,
            'phase': self.current_phase
        }
    
    def _format_time(self, seconds: float) -> str:
        """Format seconds into human-readable time"""
        if seconds is None:
            return "N/A"
        
        if seconds < 60:
            return f"{seconds:.1f}s"
        elif seconds < 3600:
            minutes = seconds / 60
            return f"{minutes:.1f}m"
        else:
            hours = seconds / 3600
            minutes = (seconds % 3600) / 60
            return f"{hours:.0f}h {minutes:.0f}m"
    
    def log_progress(self, logger, current_file: str = None, prefix: str = ""):
        """Log progress information"""
        info = self.get_progress_info(current_file)
        if not info:
            return
        
        progress_msg = (
            f"{prefix}[{info['completed']}/{info['total']}] "
            f"{info['percentage']:.1f}% | "
            f"Elapsed: {info['elapsed_time']} | "
            f"ETA: {info['eta']} | "
            f"Avg: {info['avg_time_per_file']}/file"
        )
        
        if current_file:
            progress_msg += f" | Current: {Path(current_file).name}"
        
        logger.info(progress_msg)

class DocumentProcessor:
    """Process different document types and extract text content."""
    
    @staticmethod
    def extract_text_from_pdf(file_path: Path) -> str:
        """Extract text from PDF"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.error(f"PDF extraction error for {file_path.name}: {e}")
        return text
    
    @staticmethod
    def extract_text_from_docx(file_path: Path) -> str:
        """Extract text from DOCX"""
        text = ""
        try:
            doc = docx.Document(str(file_path))
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            logger.error(f"DOCX extraction error for {file_path.name}: {e}")
        return text
    
    @staticmethod
    def extract_text_from_txt(file_path: Path) -> str:
        """Extract text from TXT with encoding detection"""
        encodings = ['utf-8', 'cp949', 'cp1252', 'latin1', 'euc-kr']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                    if text.strip():
                        return text
            except UnicodeDecodeError:
                continue
        return ""
    
    @staticmethod
    def extract_text_from_csv(file_path: Path) -> str:
        """Extract text from CSV with multiple encoding attempts"""
        encodings = ['utf-8', 'cp949', 'cp1252', 'latin1', 'euc-kr']
        
        for encoding in encodings:
            try:
                df = pd.read_csv(file_path, encoding=encoding)
                if not df.empty:
                    return df.to_string()
            except Exception:
                continue
                
        # Last resort fallback
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
                text = content.decode('utf-8', errors='replace')
                return text
        except Exception as e:
            logger.error(f"CSV extraction failed for {file_path.name}: {e}")
            
        return ""
    
    @staticmethod
    def extract_text(file_path: Path) -> str:
        """Extract text from a document based on its extension."""
        ext = file_path.suffix.lower()
        
        if ext == '.pdf':
            return DocumentProcessor.extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return DocumentProcessor.extract_text_from_docx(file_path)
        elif ext == '.txt':
            return DocumentProcessor.extract_text_from_txt(file_path)
        elif ext == '.csv':
            return DocumentProcessor.extract_text_from_csv(file_path)
        else:
            logger.warning(f"Unsupported file type: {ext}")
            return ""


class Config:
    """Configuration for the vector generator."""
    
    def __init__(self):
        # Paths
        self.input_directory = os.path.join(os.getcwd(), "batch")
        self.vector_db_path = os.path.join(os.getcwd(), "vector")
        self.plain_text_path = os.path.join(os.getcwd(), "plain")
        self.chunk_path = os.path.join(os.getcwd(), "chunk")
        
        # Chunking settings
        self.chunk_size = 1000
        self.chunk_overlap = 200
        self.max_file_size_mb = 50
        
        # Language settings
        self.languages = ["en", "ko"]
        self.language_models = {}
        
        # Embedding model settings
        self.embedding_model = "all-MiniLM-L6-v2"
        
        # Processing settings
        self.batch_size = 10
        
        # Statistics
        self.processed_files = {}
        self.processing_stats = {
            'success': 0,
            'failed': 0,
            'skipped': 0,
            'plain_documents': 0,
            'chunks': 0
        }

    def save(self, filename="vector_config.json"):
        """Save configuration to a file."""
        # Extract only serializable data
        save_data = {k: v for k, v in self.__dict__.items() 
                    if not k.startswith('_') and k != 'language_models'}
        
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(save_data, f, indent=4)
            logger.info(f"Configuration saved to {filename}")
            return True
        except Exception as e:
            logger.error(f"Failed to save configuration: {e}")
            return False
    
    def load(self, filename="vector_config.json"):
        """Load configuration from a file."""
        if os.path.exists(filename):
            try:
                with open(filename, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    for key, value in data.items():
                        if hasattr(self, key):
                            setattr(self, key, value)
                logger.info(f"Configuration loaded from {filename}")
                return True
            except Exception as e:
                logger.error(f"Failed to load configuration: {e}")
                return False
        else:
            logger.warning(f"Configuration file {filename} not found")
            return False

class VectorStoreGenerator:
    """Generate vector store from documents for RAG."""


    def __init__(self, config=None):
        self.config = config if config else Config()
        
        # Initialize paths
        self.input_directory = Path(self.config.input_directory)
        self.vector_db_path = Path(self.config.vector_db_path)
        self.plain_text_path = Path(self.config.plain_text_path)
        self.chunk_path = Path(self.config.chunk_path)
        
        # Make sure output directories exist
        self.vector_db_path.mkdir(parents=True, exist_ok=True)
        self.plain_text_path.mkdir(parents=True, exist_ok=True)
        self.chunk_path.mkdir(parents=True, exist_ok=True)
        
        # Ensure processing_stats has required keys
        if 'plain_documents' not in self.config.processing_stats:
            self.config.processing_stats['plain_documents'] = 0
        if 'chunks' not in self.config.processing_stats:
            self.config.processing_stats['chunks'] = 0
        
        # Configure chunker
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.chunk_size,
            chunk_overlap=self.config.chunk_overlap,
            length_function=len,
        )
        
        # Initialize embeddings model
        self.embeddings = HuggingFaceEmbeddings(model_name=self.config.embedding_model)
        
        # Initialize vector store
        self.vector_store = None
        self._initialize_vector_store()
        
        # Processed files tracking
        self.processed_files_path = self.vector_db_path / "processed_files.json"
        self._load_processed_files()
        
        # Supported extensions
        self.supported_extensions = ['.pdf', '.docx', '.txt', '.csv']
        
        # Progress tracker
        self.progress_tracker = ProgressTracker()
        
        # Try to load spaCy models
        self.nlp_models = {}
        self._load_language_models()
        
    def _load_language_models(self):
        """Load spaCy language models for text processing."""
        try:
            # Try to load English model
            if "en" in self.config.languages:
                try:
                    import en_core_web_sm
                    self.nlp_models["en"] = en_core_web_sm.load()
                    logger.info("Loaded English language model")
                except ImportError:
                    logger.warning("English model not found. Run: python -m spacy download en_core_web_sm")
            
            # Try to load Korean model
            if "ko" in self.config.languages:
                try:
                    import ko_core_news_sm
                    self.nlp_models["ko"] = ko_core_news_sm.load()
                    logger.info("Loaded Korean language model")
                except ImportError:
                    logger.warning("Korean model not found. Run: python -m spacy download ko_core_news_sm")
            
        except Exception as e:
            logger.error(f"Error loading language models: {e}")
    
    def _initialize_vector_store(self):
        """Initialize or load vector store."""
        vector_store_path = str(self.vector_db_path)
        
        try:
            # Check if vector store already exists
            if (Path(vector_store_path) / "index.faiss").exists():
                logger.info(f"Loading existing vector store from {vector_store_path}")
                self.vector_store = FAISS.load_local(vector_store_path, self.embeddings)
            else:
                logger.info("Initializing new vector store")
                # Will be initialized when first documents are processed
                self.vector_store = None
        except Exception as e:
            logger.error(f"Error initializing vector store: {e}")
            self.vector_store = None
    
    def _load_processed_files(self):
        """Load record of processed files"""
        if self.processed_files_path.exists():
            try:
                with open(self.processed_files_path, 'r', encoding='utf-8') as f:
                    self.config.processed_files = json.load(f)
            except Exception as e:
                logger.error(f"Error loading processed files: {e}")
                self.config.processed_files = {}
                self._save_processed_files()
        else:
            self.config.processed_files = {}
    
    def _save_processed_files(self):
        """Save record of processed files"""
        try:
            with open(self.processed_files_path, 'w', encoding='utf-8') as f:
                json.dump(self.config.processed_files, f, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.error(f"Error saving processed files: {e}")
    
    def _get_file_hash(self, file_path: Path) -> str:
        """Get MD5 hash of file"""
        import hashlib
        
        try:
            hasher = hashlib.md5()
            with open(file_path, 'rb') as f:
                buf = f.read(65536)  # Read in 64kb chunks
                while len(buf) > 0:
                    hasher.update(buf)
                    buf = f.read(65536)
            return hasher.hexdigest()
        except Exception as e:
            logger.warning(f"Cannot hash {file_path.name}: {e}")
            return str(time.time())  # Return timestamp as fallback
    
    def _get_file_size_mb(self, file_path: Path) -> float:
        """Get file size in MB"""
        try:
            size_bytes = file_path.stat().st_size
            size_mb = size_bytes / (1024 * 1024)
            return size_mb
        except Exception as e:
            logger.warning(f"Cannot get file size for {file_path}: {e}")
            return 0.0
    
    def detect_language(self, text: str) -> str:
        """Detect the language of the text."""
        if not text or len(text.strip()) < 10:
            return "en"  # Default to English for empty or very short texts
        
        # Simple heuristic: count Korean and English characters
        ko_char_count = len(re.findall(r'[가-힣]', text))
        en_char_count = len(re.findall(r'[a-zA-Z]', text))
        
        if ko_char_count > en_char_count:
            return "ko"
        else:
            return "en"
    
    def clean_text(self, text: str, language: str = "en") -> str:
        """Clean and normalize text based on detected language."""
        if not text:
            return ""
        
        # Basic cleaning
        text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
        text = text.strip()
        
        # Use spaCy for advanced processing if model is available
        if language in self.nlp_models:
            nlp = self.nlp_models[language]
            # Process with spaCy, but limit document size to avoid memory issues
            max_length = 100000  # 100k characters max
            if len(text) > max_length:
                logger.warning(f"Text too long for spaCy processing, truncating to {max_length} characters")
                text = text[:max_length]
            
            try:
                doc = nlp(text)
                # Basic sentence normalization
                processed_text = " ".join([sent.text.strip() for sent in doc.sents])
                return processed_text
            except Exception as e:
                logger.error(f"Error in spaCy processing: {e}")
                return text
        
        return text

    def process_file(self, file_path: Path) -> bool:
        """Process a single file and add it to the vector store."""
        try:
            start_time = time.time()
            file_str = str(file_path)
            
            # Check if file exists
            if not file_path.exists():
                logger.warning(f"File doesn't exist: {file_path}")
                self.config.processing_stats['skipped'] += 1
                return False
            
            # Check file size
            file_size = file_path.stat().st_size
            size_mb = self._get_file_size_mb(file_path)
            
            logger.debug(f"File {file_path.name}: {file_size} bytes ({size_mb:.3f} MB)")
            
            if file_size == 0:
                logger.warning(f"Skipping empty file: {file_path.name}")
                self.config.processing_stats['skipped'] += 1
                return False
            
            if size_mb > self.config.max_file_size_mb:
                logger.warning(f"Skipping large file: {file_path.name} ({size_mb:.3f} MB)")
                self.config.processing_stats['skipped'] += 1
                return False
            
            # Get file hash for change detection
            file_hash = self._get_file_hash(file_path)
            
            # Check if already processed and unchanged
            if file_str in self.config.processed_files:
                if self.config.processed_files[file_str]['hash'] == file_hash:
                    logger.info(f"Already processed and unchanged: {file_path.name}")
                    return True
            
            logger.info(f"Processing: {file_path.name} ({size_mb:.3f} MB)")
            
            # Extract text
            text = DocumentProcessor.extract_text(file_path)
            if not text or not text.strip():
                logger.warning(f"No text extracted from {file_path.name}")
                self.config.processing_stats['failed'] += 1
                return False
            
            # Detect language
            language = self.detect_language(text)
            logger.info(f"Detected language for {file_path.name}: {language}")
            
            # Clean text
            text = self.clean_text(text, language)
            
            # Ensure plain_documents key exists in processing_stats
            if 'plain_documents' not in self.config.processing_stats:
                self.config.processing_stats['plain_documents'] = 0
            
            # Save plain text file
            plain_file_path = self.plain_text_path / f"{file_path.stem}_{file_hash[:8]}.txt"
            try:
                with open(plain_file_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                logger.info(f"Saved plain text to {plain_file_path}")
                self.config.processing_stats['plain_documents'] += 1
            except Exception as e:
                logger.error(f"Failed to save plain text file: {e}")
                # Continue processing even if plain text save fails
            
            # Split into chunks
            chunks = self.text_splitter.split_text(text)
            if not chunks:
                logger.warning(f"No chunks created for {file_path.name}")
                self.config.processing_stats['failed'] += 1
                return False
            
            # Ensure chunks key exists in processing_stats
            if 'chunks' not in self.config.processing_stats:
                self.config.processing_stats['chunks'] = 0
                
            # Create chunks directory with safer name (avoid special characters)
            safe_stem = re.sub(r'[^\w\-.]', '_', file_path.stem)
            chunks_directory = self.chunk_path / f"{safe_stem}_{file_hash[:8]}"
            chunks_directory.mkdir(exist_ok=True, parents=True)
            
            # Log chunk creation
            logger.info(f"Created chunk directory: {chunks_directory}")
            
            for i, chunk in enumerate(chunks):
                # Log the first few lines of chunk for verification (line 3)
                chunk_lines = chunk.split('\n')
                if i < 3:  # Log only first 3 chunks
                    if len(chunk_lines) >= 3:
                        logger.info(f"Chunk {i} - Line 3: {chunk_lines[2][:50]}...")
                    elif len(chunk_lines) > 0:
                        logger.info(f"Chunk {i} - First line: {chunk_lines[0][:50]}...")
                
                # Save chunk to file
                chunk_file_path = chunks_directory / f"chunk_{i:04d}.txt"
                try:
                    with open(chunk_file_path, 'w', encoding='utf-8') as f:
                        f.write(chunk)
                except Exception as e:
                    logger.error(f"Failed to save chunk {i} to file: {e}")
                    # Continue processing even if chunk save fails
            
            # Update chunks count
            self.config.processing_stats['chunks'] += len(chunks)
            
            # Create document metadata
            metadata = {
                'source': file_str,
                'filename': file_path.name,
                'file_hash': file_hash,
                'language': language,
                'processed_date': datetime.now().isoformat(),
                'chunks': len(chunks),
                'plain_text_path': str(plain_file_path),
                'chunks_directory': str(chunks_directory)
            }
            
            # Process chunks for vector store
            documents = []
            for i, chunk in enumerate(chunks):
                doc_metadata = metadata.copy()
                doc_metadata['chunk_index'] = i
                doc_metadata['chunk_file'] = str(chunks_directory / f"chunk_{i:04d}.txt")
                documents.append(Document(page_content=chunk, metadata=doc_metadata))
            
            # Add to vector store
            if documents:
                try:
                    if not self.vector_store:
                        self.vector_store = FAISS.from_documents(documents, self.embeddings)
                        logger.info(f"Created new vector store with {len(documents)} document chunks")
                    else:
                        self.vector_store.add_documents(documents)
                        logger.info(f"Added {len(documents)} chunks to vector store")
                    
                    # Save immediately
                    self._save_vector_store()
                except Exception as e:
                    logger.error(f"Vector store update failed for {file_path.name}: {e}")
                    self.config.processing_stats['failed'] += 1
                    return False
            
            # Update processed files
            self.config.processed_files[file_str] = {
                'hash': file_hash,
                'language': language,
                'processed_date': datetime.now().isoformat(),
                'chunks': len(chunks),
                'processing_time': time.time() - start_time,
                'plain_text_path': str(plain_file_path),
                'chunks_directory': str(chunks_directory)
            }
            self.config.processing_stats['success'] += 1
            self._save_processed_files()
            
            processing_time = time.time() - start_time
            logger.info(f"Successfully processed {file_path.name}: {len(chunks)} chunks in {processing_time:.2f}s")
            return True
            
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            logger.error(traceback.format_exc())
            self.config.processing_stats['failed'] += 1
            return False
                

    def _save_vector_store(self):
        """Save vector store to disk"""
        try:
            if self.vector_store:
                self.vector_store.save_local(str(self.vector_db_path))
                logger.info(f"Vector store saved to {self.vector_db_path}")
        except Exception as e:
            logger.error(f"Failed to save vector store: {e}")
    
    def index_directory(self, batch_size: int = None):
        """Index all documents in the directory."""
        if batch_size is None:
            batch_size = self.config.batch_size
            
        if not self.input_directory.exists():
            logger.error(f"Directory not found: {self.input_directory}")
            return
        
        # Find all supported files
        files_to_process = []
        for ext in self.supported_extensions:
            files_to_process.extend(self.input_directory.rglob(f"*{ext}"))
        
        # Filter and sort files
        valid_files = []
        for file_path in files_to_process:
            try:
                if file_path.exists() and file_path.is_file():
                    valid_files.append(file_path)
            except Exception as e:
                logger.warning(f"Cannot access file {file_path}: {e}")
                continue
        
        # Sort files by size (smaller files first)
        valid_files.sort(key=lambda f: self._get_file_size_mb(f) if f.exists() else 0)
        
        total_files = len(valid_files)
        logger.info(f"Found {total_files} files to process in {self.input_directory}")
        
        if total_files == 0:
            logger.warning("No files found to index")
            return
        
        # Initialize progress tracker
        self.progress_tracker.start_tracking(total_files, "indexing")
        
        # Reset processing stats
        self.config.processing_stats = {
            'success': 0,
            'failed': 0,
            'skipped': 0
        }
        
        # Process each file
        for i, file_path in enumerate(valid_files, 1):
            try:
                self.progress_tracker.start_file()
                
                # Process file
                success = self.process_file(file_path)
                
                # Update progress
                self.progress_tracker.complete_file(str(file_path))
                
                # Log progress periodically
                if i % 5 == 0 or i == total_files:
                    info = self.progress_tracker.get_progress_info()
                    logger.info(
                        f"Progress: {info['completed']}/{info['total']} files "
                        f"({info['percentage']:.1f}%) | "
                        f"Success: {self.config.processing_stats['success']}, "
                        f"Failed: {self.config.processing_stats['failed']}, "
                        f"Skipped: {self.config.processing_stats['skipped']}"
                    )
                
                # Save progress periodically
                if i % batch_size == 0:
                    try:
                        self._save_vector_store()
                        self._save_processed_files()
                        logger.info(f"Batch saved at {i} files")
                    except Exception as e:
                        logger.error(f"Error saving batch: {e}")
            
            except Exception as e:
                logger.error(f"Error processing file {i}/{total_files} ({file_path}): {e}")
                self.progress_tracker.complete_file(str(file_path))
                continue
        
        # Final save
        try:
            self._save_vector_store()
            self._save_processed_files()
        except Exception as e:
            logger.error(f"Error in final save: {e}")
        
        final_info = self.progress_tracker.get_progress_info()
        logger.info(
            f"Indexing completed: {final_info['completed']}/{final_info['total']} files "
            f"in {final_info['elapsed_time']}"
        )
        self._log_stats()

    def get_statistics(self) -> Dict:
        """Get indexing statistics"""
        vector_files = list(self.vector_db_path.glob("*.faiss"))
        vector_index_exists = any(vector_files)
        
        # Count files in directories
        plain_files_count = len(list(self.plain_text_path.glob("*.txt")))
        chunk_dirs_count = len([d for d in self.chunk_path.iterdir() if d.is_dir()])
        
        # Ensure processing_stats has required keys
        if 'plain_documents' not in self.config.processing_stats:
            self.config.processing_stats['plain_documents'] = 0
        if 'chunks' not in self.config.processing_stats:
            self.config.processing_stats['chunks'] = 0
        
        stats = {
            'total_processed': len(self.config.processed_files),
            'processing_stats': self.config.processing_stats.copy(),
            'vector_store_exists': vector_index_exists,
            'vector_db_path': str(self.vector_db_path),
            'input_directory': str(self.input_directory),
            'plain_text_path': str(self.plain_text_path),
            'chunk_path': str(self.chunk_path),
            'plain_files_count': plain_files_count,
            'chunk_dirs_count': chunk_dirs_count,
            'embedding_model': self.config.embedding_model,
            'chunk_size': self.config.chunk_size,
            'chunk_overlap': self.config.chunk_overlap,
            'languages': list(self.nlp_models.keys())
        }
        
        # Get language distribution if documents exist
        if self.config.processed_files:
            language_stats = {}
            for file_info in self.config.processed_files.values():
                lang = file_info.get('language', 'unknown')
                if lang in language_stats:
                    language_stats[lang] += 1
                else:
                    language_stats[lang] = 1
            stats['language_distribution'] = language_stats
        
        # Get total chunks count
        total_chunks = sum(
            file_info.get('chunks', 0) 
            for file_info in self.config.processed_files.values()
        )
        stats['total_chunks'] = total_chunks
        
        return stats

    def _log_stats(self):
        """Log current processing statistics"""
        logger.info(f"Processing stats - Success: {self.config.processing_stats['success']}, "
                   f"Failed: {self.config.processing_stats['failed']}, "
                   f"Skipped: {self.config.processing_stats['skipped']}")

class VectorStoreApp:
    """Main application class for the vector store generator."""
    
    def __init__(self):
        self.config = Config()
        
        # Try to load configuration
        try:
            self.config.load()
        except:
            # If loading fails, use default configuration
            pass
        
        # Ensure directories exist
        os.makedirs(self.config.input_directory, exist_ok=True)
        os.makedirs(self.config.vector_db_path, exist_ok=True)
        
        # Initialize vector store generator
        self.vector_generator = VectorStoreGenerator(self.config)
    
    def scan_directory(self, custom_dir: str = None):
        """Scan and index directory."""
        if custom_dir:
            orig_input_dir = self.config.input_directory
            self.config.input_directory = custom_dir
            self.vector_generator.input_directory = Path(custom_dir)
        
        try:
            # Start indexing
            self.vector_generator.index_directory()
            
            # Get and display statistics
            stats = self.vector_generator.get_statistics()
            
            console.print("\n[bold green]Indexing completed![/bold green]")
            console.print(f"[bold]Total Files Processed:[/bold] {stats['total_processed']}")
            console.print(f"[bold]Total Chunks Created:[/bold] {stats['total_chunks']}")
            console.print(f"[bold]Success:[/bold] {stats['processing_stats']['success']}")
            console.print(f"[bold]Failed:[/bold] {stats['processing_stats']['failed']}")
            console.print(f"[bold]Skipped:[/bold] {stats['processing_stats']['skipped']}")
            
            if 'language_distribution' in stats:
                console.print("\n[bold]Language Distribution:[/bold]")
                for lang, count in stats['language_distribution'].items():
                    console.print(f"  - {lang}: {count} files")
            
        finally:
            # Restore original directory if using custom dir
            if custom_dir:
                self.config.input_directory = orig_input_dir
                self.vector_generator.input_directory = Path(orig_input_dir)
    
    def show_statistics(self):
        """Show application statistics."""
        stats = self.vector_generator.get_statistics()
        
        console.print("[bold blue]Vector Store Statistics[/bold blue]")

        console.print(f"\n[bold]General Information:[/bold]")
        console.print(f"Input Directory: {stats['input_directory']}")
        console.print(f"Plain Text Path: {stats['plain_text_path']}")
        console.print(f"Chunk Path: {stats['chunk_path']}")
        console.print(f"Vector Store Path: {stats['vector_db_path']}")
        console.print(f"Vector Store Exists: {'Yes' if stats['vector_store_exists'] else 'No'}")
        console.print(f"Embedding Model: {stats['embedding_model']}")
        
        console.print(f"\n[bold]Configuration:[/bold]")
        console.print(f"Chunk Size: {stats['chunk_size']}")
        console.print(f"Chunk Overlap: {stats['chunk_overlap']}")
        console.print(f"Supported Languages: {', '.join(stats['languages'])}")
        
        console.print(f"\n[bold]Processing Statistics:[/bold]")
        console.print(f"Files Processed: {stats['total_processed']}")
        console.print(f"Plain Documents: {stats['plain_files_count']}")
        console.print(f"Chunk Directories: {stats['chunk_dirs_count']}")
        console.print(f"Document Chunks: {stats['total_chunks']}")
        
        proc_stats = stats['processing_stats']
        console.print(f"Success: {proc_stats['success']}")
        console.print(f"Failed: {proc_stats['failed']}")
        console.print(f"Skipped: {proc_stats['skipped']}")
        
        # If we have language stats, show them
        if 'language_distribution' in stats and stats['language_distribution']:
            console.print(f"\n[bold]Language Distribution:[/bold]")
            
            lang_table = Table()
            lang_table.add_column("Language", style="cyan")
            lang_table.add_column("Files", style="green")
            lang_table.add_column("Percentage", style="yellow")
            
            total = sum(stats['language_distribution'].values())
            
            for lang, count in sorted(stats['language_distribution'].items(), 
                                    key=lambda x: x[1], reverse=True):
                percentage = (count / total) * 100
                lang_table.add_row(
                    lang,
                    str(count),
                    f"{percentage:.1f}%"
                )
            
            console.print(lang_table)
                
    def configure_app(self):
        """Configure application settings."""
        console.print("[bold blue]Application Configuration[/bold blue]")
        
        # Input/output directories
        console.print("\n[bold]Directory Settings:[/bold]")
        
        self.config.input_directory = Prompt.ask(
            "Input directory path (batch)",
            default=self.config.input_directory
        )
        
        self.config.plain_text_path = Prompt.ask(
            "Plain text directory path",
            default=self.config.plain_text_path
        )
        
        self.config.chunk_path = Prompt.ask(
            "Chunk directory path",
            default=self.config.chunk_path
        )
        
        self.config.vector_db_path = Prompt.ask(
            "Vector database path",
            default=self.config.vector_db_path
        )
        
        # Chunking settings
        console.print("\n[bold]Text Chunking Settings:[/bold]")
        
        self.config.chunk_size = int(Prompt.ask(
            "Chunk size (characters)",
            default=str(self.config.chunk_size)
        ))
        
        self.config.chunk_overlap = int(Prompt.ask(
            "Chunk overlap (characters)",
            default=str(self.config.chunk_overlap)
        ))
        
        self.config.max_file_size_mb = int(Prompt.ask(
            "Maximum file size (MB)",
            default=str(self.config.max_file_size_mb)
        ))
        
        # Language settings
        console.print("\n[bold]Language Settings:[/bold]")
        console.print("Select languages to support:")
        console.print("1. English only")
        console.print("2. Korean only")
        console.print("3. Both English and Korean")
        
        lang_choice = Prompt.ask(
            "Choose language option",
            choices=["1", "2", "3"],
            default="3"
        )
        
        if lang_choice == "1":
            self.config.languages = ["en"]
        elif lang_choice == "2":
            self.config.languages = ["ko"]
        else:
            self.config.languages = ["en", "ko"]
        
        # Embedding model
        console.print("\n[bold]Embedding Model:[/bold]")
        console.print("1. all-MiniLM-L6-v2 (fast, lower quality)")
        console.print("2. msmarco-distilbert-base-v4 (balanced)")
        console.print("3. all-mpnet-base-v2 (slower, high quality)")
        
        model_choice = Prompt.ask(
            "Choose embedding model",
            choices=["1", "2", "3"],
            default="1"
        )
        
        if model_choice == "1":
            self.config.embedding_model = "all-MiniLM-L6-v2"
        elif model_choice == "2":
            self.config.embedding_model = "msmarco-distilbert-base-v4"
        elif model_choice == "3":
            self.config.embedding_model = "all-mpnet-base-v2"
        
        # Batch settings
        self.config.batch_size = int(Prompt.ask(
            "Batch size (save frequency)",
            default=str(self.config.batch_size)
        ))
        
        # Save configuration
        if self.config.save():
            console.print("[green]Configuration saved successfully![/green]")
        else:
            console.print("[red]Failed to save configuration.[/red]")
        
        # Ask to reinitialize vector generator with new settings
        if Confirm.ask("Apply new settings now?", default=True):
            self.vector_generator = VectorStoreGenerator(self.config)
            console.print("[green]Settings applied![/green]")
            

    def test_query(self):
        """Test querying the vector store."""
        # Check if vector store exists and is initialized
        if not self.vector_generator.vector_store:
            console.print("[red]Vector store not initialized or empty. Index some documents first.[/red]")
            return
        
        console.print("[bold blue]Test Vector Store Query[/bold blue]")
        console.print("Enter a query to test retrieval from the vector store.")
        
        # Get query from user
        query = Prompt.ask("Query")
        
        if not query.strip():
            console.print("[yellow]Empty query. Please enter a valid query.[/yellow]")
            return
        
        # Get number of results to show
        k = int(Prompt.ask("Number of results to show", default="5"))
        
        console.print(f"\n[bold]Searching for:[/bold] {query}")
        console.print("[bold]Results:[/bold]")
        
        try:
            # Perform similarity search
            with Progress(
                SpinnerColumn(),
                TextColumn("[bold blue]Searching vector store..."),
                console=console
            ) as progress:
                task = progress.add_task("Searching", total=1)
                results = self.vector_generator.vector_store.similarity_search_with_score(query, k=k)
                progress.update(task, completed=1)
            
            if not results:
                console.print("[yellow]No results found.[/yellow]")
                return
            
            # Display results
            result_table = Table(title=f"Top {len(results)} Results")
            result_table.add_column("Score", style="cyan")
            result_table.add_column("Source", style="green")
            result_table.add_column("Chunk", style="magenta")
            result_table.add_column("Content", style="white")
            
            for i, (doc, score) in enumerate(results, 1):
                # Format content text for display (truncate if too long)
                content = doc.page_content.strip()
                if len(content) > 200:
                    content = content[:197] + "..."
                
                # Format score (lower is better for cosine distance)
                formatted_score = f"{1.0 - score:.4f}"
                
                # Get metadata
                source = Path(doc.metadata.get('source', 'Unknown')).name
                chunk_index = doc.metadata.get('chunk_index', 'Unknown')
                
                result_table.add_row(
                    formatted_score,
                    source,
                    str(chunk_index),
                    content
                )
            
            console.print(result_table)
            
            # Ask if user wants to see detailed content
            if Confirm.ask("Show full content of a specific result?", default=False):
                result_num = int(Prompt.ask(
                    f"Enter result number (1-{len(results)})",
                    default="1"
                ))
                
                if 1 <= result_num <= len(results):
                    doc, score = results[result_num - 1]
                    
                    # Display full content in a panel
                    metadata = doc.metadata
                    metadata_str = "\n".join([f"{k}: {v}" for k, v in metadata.items()])
                    
                    console.print(Panel(
                        f"[bold]Metadata:[/bold]\n{metadata_str}\n\n[bold]Content:[/bold]\n{doc.page_content}",
                        title=f"Result {result_num} Details",
                        expand=False
                    ))
                else:
                    console.print("[yellow]Invalid result number.[/yellow]")
        
        except Exception as e:
            console.print(f"[red]Error during search: {str(e)}[/red]")
            logger.error(f"Error during vector search: {e}")
            logger.error(traceback.format_exc())
    
    def clear_vector_store(self):
        """Clear the vector store."""
        console.print("[bold red]Clear Vector Store[/bold red]")
        console.print("This will delete the entire vector store and reset processing history.")
        
        if not Confirm.ask("Are you sure you want to continue?", default=False):
            console.print("[yellow]Operation cancelled.[/yellow]")
            return
        
        try:
            # Delete vector store files
            vector_path = Path(self.config.vector_db_path)
            
            # Find all vector store related files
            vector_files = list(vector_path.glob("*.faiss")) + list(vector_path.glob("*.pkl"))
            
            if vector_files:
                # Delete each file
                for file in vector_files:
                    file.unlink()
                
                # Also delete processed files record
                processed_files_path = vector_path / "processed_files.json"
                if processed_files_path.exists():
                    processed_files_path.unlink()
                
                # Reset processing records
                self.config.processed_files = {}
                self.config.processing_stats = {
                    'success': 0,
                    'failed': 0,
                    'skipped': 0
                }
                
                # Reinitialize vector generator
                self.vector_generator = VectorStoreGenerator(self.config)
                
                console.print("[green]Vector store cleared successfully![/green]")
            else:
                console.print("[yellow]No vector store files found to delete.[/yellow]")
        
        except Exception as e:
            console.print(f"[red]Error clearing vector store: {str(e)}[/red]")
            logger.error(f"Error clearing vector store: {e}")

    def run(self):
        """Run the CLI application."""
        while True:
            console.clear()
            console.print("[bold blue]Vector Store Generator for RAG[/bold blue]")
            console.print("[yellow]=============================[/yellow]")
            console.print("A tool to process documents for vector-based retrieval")
            
            console.print("\n[bold]Current Configuration:[/bold]")
            console.print(f"Input Directory (batch): {self.config.input_directory}")
            console.print(f"Plain Text Path: {self.config.plain_text_path}")
            console.print(f"Chunk Path: {self.config.chunk_path}")
            console.print(f"Vector Store Path: {self.config.vector_db_path}")
            console.print(f"Embedding Model: {self.config.embedding_model}")
            console.print(f"Chunk Size: {self.config.chunk_size}")
            console.print(f"Languages: {', '.join(self.config.languages)}")
            
            # Create menu
            console.print("\n[bold]Menu:[/bold]")
            console.print("1. Process Documents")
            console.print("2. Test Query")
            console.print("3. Show Statistics")
            console.print("4. Configure Application")
            console.print("5. Clear Vector Store")
            console.print("0. Exit")
            
            choice = Prompt.ask("Select an option", choices=["0", "1", "2", "3", "4", "5"], default="1")
        
            
            if choice == "0":
                break
            elif choice == "1":
                # Process documents
                custom_dir = Confirm.ask("Use custom input directory?", default=False)
                dir_path = None
                
                if custom_dir:
                    dir_path = Prompt.ask("Enter directory path", default=self.config.input_directory)
                
                self.scan_directory(dir_path)
                Prompt.ask("Press Enter to continue")
            elif choice == "2":
                # Test query
                self.test_query()
                Prompt.ask("Press Enter to continue")
            elif choice == "3":
                # Show statistics
                self.show_statistics()
                Prompt.ask("Press Enter to continue")
            elif choice == "4":
                # Configure application
                self.configure_app()
                Prompt.ask("Press Enter to continue")
            elif choice == "5":
                # Clear vector store
                self.clear_vector_store()
                Prompt.ask("Press Enter to continue")
        
        console.print("[green]Thank you for using Vector Store Generator![/green]")

def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(description="Vector Store Generator for RAG")
    parser.add_argument("--input-dir", help="Input directory containing documents")
    parser.add_argument("--vector-dir", help="Output directory for vector store")
    parser.add_argument("--chunk-size", type=int, help="Text chunk size")
    parser.add_argument("--chunk-overlap", type=int, help="Text chunk overlap")
    parser.add_argument("--max-file-size", type=int, help="Maximum file size in MB")
    parser.add_argument("--batch-size", type=int, help="Batch size for processing")
    parser.add_argument("--model", help="Embedding model name")
    parser.add_argument("--interactive", action="store_true", help="Run in interactive mode")
    
    args = parser.parse_args()
    
    try:
        # Create app
        app = VectorStoreApp()
        
        # If CLI arguments are provided, apply them
        if args.input_dir:
            app.config.input_directory = args.input_dir
        if args.vector_dir:
            app.config.vector_db_path = args.vector_dir
        if args.chunk_size:
            app.config.chunk_size = args.chunk_size
        if args.chunk_overlap:
            app.config.chunk_overlap = args.chunk_overlap
        if args.max_file_size:
            app.config.max_file_size_mb = args.max_file_size
        if args.batch_size:
            app.config.batch_size = args.batch_size
        if args.model:
            app.config.embedding_model = args.model
        
        # If CLI arguments are provided and not interactive, update generator and process
        if any([args.input_dir, args.vector_dir, args.chunk_size, args.chunk_overlap, 
               args.max_file_size, args.batch_size, args.model]) and not args.interactive:
            app.vector_generator = VectorStoreGenerator(app.config)
            app.scan_directory()
            return
        
        # Run interactive app if requested or by default
        if args.interactive or len(sys.argv) == 1:
            app.run()
        
    except KeyboardInterrupt:
        console.print("\n[yellow]Application terminated by user.[/yellow]")
    except Exception as e:
        console.print(f"[red]An unexpected error occurred: {str(e)}[/red]")
        logger.error(f"Unexpected error: {str(e)}\n{traceback.format_exc()}")
    finally:
        console.print("[green]Exiting application. Goodbye![/green]")

if __name__ == "__main__":
    main()